from docx import Document

# Function to create the document with specific content
def create_statement_doc(filename, content, title):
    doc = Document()
    doc.add_heading("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", level=1)
    doc.add_heading("Độc lập - Tự do - Hạnh phúc", level=1)
    doc.add_paragraph("\n")
    doc.add_heading(title, level=2)
    doc.add_paragraph("\n")
    for line in content:
        doc.add_paragraph(line)
    doc.save(filename)

# Content for each document

# 1. Tờ trình gửi cơ quan thanh tra công an để làm rõ sự việc
contpent_1 = [
    "Kính gửi: Cơ quan Thanh tra - Công an quận Nam Từ Liêm",
    "Tôi tên là: Nguyễn Thị Thu Hương",
    "Ngày tháng năm sinh: 11/04/2000",
    "CMND/CCCD: 064300008177",
    "Địa chỉ thường trú: Thôn Thái Hà, Ia Yok, Ia Grai, Gia Lai",
    "Điện thoại liên hệ: +84352987654",
    "",
    "Nội dung trình bày:",
    "Tôi xin trình bày rõ sự việc xảy ra liên quan đến việc tôi bị triệu tập và lời khai của tôi tại Công an quận Nam Từ Liêm. ",
    "Tôi khẳng định rằng tôi không liên quan đến bất kỳ hành vi vi phạm pháp luật nào và tôi yêu cầu cơ quan thanh tra làm rõ các sự việc liên quan, đặc biệt là:",
    "1. Việc triệu tập và nội dung lời khai của tôi tại Công an Nam Từ Liêm.",
    "2. Kiểm tra các tài liệu, hồ sơ, để đảm bảo không có sai sót hoặc việc ghi khống thông tin.",
    "Tôi mong muốn được nhận sự hỗ trợ minh bạch và công tâm từ cơ quan thanh tra để làm rõ những vấn đề nêu trên.",
    "",
    "Tôi xin chân thành cảm ơn!",
    "",
    "Người làm đơn,",
    "(Ký và ghi rõ họ tên)"
]

# 2. Tờ trình yêu cầu Công an Nam Từ Liêm cung cấp bản sao hồ sơ lấy lời khai
content_2 = [
    "Kính gửi: Cơ quan Cảnh sát điều tra - Công an quận Nam Từ Liêm",
    "Tôi tên là: Nguyễn Thị Thu Hương",
    "Ngày tháng năm sinh: 11/04/2000",
    "CMND/CCCD: 064300008177",
    "Địa chỉ thường trú: Thôn Thái Hà, Ia Yok, Ia Grai, Gia Lai",
    "Điện thoại liên hệ: +84352987654",
    "",
    "Nội dung trình bày:",
    "Tôi xin đề nghị quý cơ quan cung cấp bản sao hồ sơ lời khai của tôi được lập tại Công an quận Nam Từ Liêm. ",
    "Lý do: Tôi cần kiểm tra và đối chiếu để đảm bảo rằng nội dung lời khai không bị chỉnh sửa, ghi khống hoặc có sai sót.",
    "Tôi đề nghị quý cơ quan hỗ trợ và tạo điều kiện để tôi nhận được bản sao này trong thời gian sớm nhất.",
    "",
    "Tôi xin chân thành cảm ơn!",
    "",
    "Người làm đơn,",
    "(Ký và ghi rõ họ tên)"
]

# 3. Tờ trình yêu cầu Viện Kiểm sát cung cấp hồ sơ tố tụng
content_3 = [
    "Kính gửi: Viện Kiểm sát nhân dân quận Nam Từ Liêm",
    "Tôi tên là: Nguyễn Thị Thu Hương",
    "Ngày tháng năm sinh: 11/04/2000",
    "CMND/CCCD: 064300008177",
    "Địa chỉ thường trú: Thôn Thái Hà, Ia Yok, Ia Grai, Gia Lai",
    "Điện thoại liên hệ: +84352987654",
    "",
    "Nội dung trình bày:",
    "Tôi xin đề nghị quý Viện Kiểm sát cung cấp hồ sơ tố tụng liên quan đến vụ việc tôi bị triệu tập.",
    "Lý do: Đến nay, tôi chưa ký vào bất kỳ hồ sơ tố tụng nào và cần xác minh các tài liệu có liên quan.",
    "Tôi hy vọng quý cơ quan sẽ hỗ trợ để tôi có thể nhận được hồ sơ này trong thời gian sớm nhất.",
    "",
    "Tôi xin chân thành cảm ơn!",
    "",
    "Người làm đơn,",
    "(Ký và ghi rõ họ tên)"
]

# File paths
file_1 = "/mnt/data/To_Trinh_Thanh_Tra_Cong_An.docx"
file_2 = "/mnt/data/To_Trinh_Cong_An_Nam_Tu_Liem.docx"
file_3 = "/mnt/data/To_Trinh_Vien_Kiem_Sat.docx"

# Create the documents
create_statement_doc(file_1, content_1, "TỜ TRÌNH VỀ VIỆC LÀM RÕ SỰ VIỆC")
create_statement_doc(file_2, content_2, "TỜ TRÌNH VỀ VIỆC CUNG CẤP HỒ SƠ LẤY LỜI KHAI")
create_statement_doc(file_3, content_3, "TỜ TRÌNH VỀ VIỆC CUNG CẤP HỒ SƠ TỐ TỤNG")

file_1, file_2, file_3
